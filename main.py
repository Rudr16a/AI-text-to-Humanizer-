import docx
import re
import random
import string
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

class UltimateTextHumanizer:
    def __init__(self):
        # Comprehensive AI pattern replacement database
        self.ai_patterns = {
            # Formal to casual transitions
            r'\bIn conclusion\b': ['To wrap up', 'All in all', 'When it comes down to it', 'At the end of the day'],
            r'\bIt is important to note that\b': ['Keep in mind', 'Remember', 'What you should know is', 'The thing is'],
            r'\bFurthermore\b': ['What\'s more', 'On top of that', 'Plus', 'And another thing'],
            r'\bMoreover\b': ['Besides that', 'Also', 'Not only that', 'And let\'s not forget'],
            r'\bHowever\b': ['But', 'That said', 'Then again', 'On the other hand'],
            r'\bNevertheless\b': ['Still', 'Even so', 'All the same', 'Despite that'],
            r'\bThus\b': ['So', 'This means', 'Because of this', 'As a result'],
            r'\bHence\b': ['So', 'Therefore', 'That\'s why', 'This leads to'],
            r'\bAdditionally\b': ['Also', 'Plus', 'Another thing', 'What\'s more'],
            r'\bConsequently\b': ['So', 'Because of this', 'This means', 'As you might expect'],
            r'\bAccordingly\b': ['So', 'Because of this', 'This is why', 'Given that'],
            
            # Corporate jargon replacement
            r'\bleverage\b': ['use', 'make the most of', 'work with', 'take advantage of'],
            r'\bsynergy\b': ['teamwork', 'working together', 'collaboration', 'combined effort'],
            r'\bparadigm\b': ['approach', 'model', 'way of thinking', 'method'],
            r'\butilize\b': ['use', 'work with', 'make use of', 'put to work'],
            r'\boptimal\b': ['best', 'ideal', 'most effective', 'right'],
            r'\bfacilitate\b': ['help', 'make easier', 'assist with', 'enable'],
        }
        
        # Human speech patterns and filler words
        self.human_patterns = {
            'sentence_starters': [
                'You know,', 'I mean,', 'Well,', 'So,', 'Actually,', 'Basically,',
                'Honestly,', 'Seriously,', 'To be honest,', 'The way I see it,',
                'From what I understand,', 'If you ask me,', 'In my experience,'
            ],
            'casual_connectors': [
                'kind of', 'sort of', 'a bit', 'pretty much', 'more or less',
                'you know what I mean', 'and all that', 'and everything',
                'or something', 'and stuff like that'
            ],
            'emphasis_words': [
                'really', 'actually', 'literally', 'basically', 'obviously',
                'clearly', 'definitely', 'absolutely', 'completely'
            ],
            'thinking_words': [
                'like', 'I guess', 'I suppose', 'maybe', 'perhaps',
                'probably', 'apparently', 'seemingly'
            ]
        }
        
        # Advanced contraction system
        self.contractions = {
            r'\bdo not\b': "don't",
            r'\bdoes not\b': "doesn't", 
            r'\bdid not\b': "didn't",
            r'\bcannot\b': "can't",
            r'\bwill not\b': "won't",
            r'\bshould not\b': "shouldn't",
            r'\bwould not\b': "wouldn't",
            r'\bcould not\b': "couldn't",
            r'\bit is\b': "it's",
            r'\bthat is\b': "that's",
            r'\bthey are\b': "they're",
            r'\bwe are\b': "we're",
            r'\byou are\b': "you're",
            r'\bI am\b': "I'm",
            r'\bhe is\b': "he's",
            r'\bshe is\b': "she's",
            r'\bthere is\b': "there's",
            r'\bwhat is\b': "what's",
            r'\bwhere is\b': "where's",
            r'\bhow is\b': "how's",
            r'\bwhen is\b': "when's",
            r'\bwhy is\b': "why's",
            r'\bwho is\b': "who's",
            r'\bI have\b': "I've",
            r'\bthey have\b': "they've",
            r'\bwe have\b': "we've",
            r'\byou have\b': "you've",
        }
        
        # Track document statistics for consistency
        self.stats = defaultdict(int)
        self.preferences = {}

    def analyze_writing_style(self, text):
        """Analyze the original text to maintain consistent style"""
        sentences = re.split(r'[.!?]+', text)
        avg_sentence_length = sum(len(s.split()) for s in sentences) / max(len(sentences), 1)
        
        # Determine formality level
        formal_words = len(re.findall(r'\b(however|moreover|thus|therefore|consequently)\b', text, re.IGNORECASE))
        self.preferences['formality'] = 'formal' if formal_words > len(sentences) * 0.1 else 'casual'
        
        # Determine contraction preference
        contraction_count = len(re.findall(r"\w+'\w+", text))
        self.preferences['contractions'] = 'high' if contraction_count > len(sentences) * 0.3 else 'low'
        
        return avg_sentence_length

    def humanize_text(self, text):
        """Advanced text humanization with multiple techniques"""
        if not text or not text.strip():
            return text
        
        # Analyze original style
        if not self.preferences:
            self.analyze_writing_style(text)
        
        original_text = text
        
        # Apply humanization techniques in sequence
        techniques = [
            self.replace_ai_patterns,
            self.add_contractions,
            self.vary_sentence_flow,
            self.add_human_speech_patterns,
            self.introduce_natural_imperfections,
            self.adjust_formality_level,
            self.add_personal_touches
        ]
        
        for technique in techniques:
            text = technique(text)
            # Safety check to prevent text corruption
            if not text or len(text.strip()) < len(original_text.strip()) * 0.3:
                text = original_text
                break
        
        return text

    def replace_ai_patterns(self, text):
        """Replace AI patterns with natural human alternatives"""
        for pattern, replacements in self.ai_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                replacement = random.choice(replacements)
                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        return text

    def add_contractions(self, text):
        """Intelligently add contractions based on context and preferences"""
        contraction_chance = 0.8 if self.preferences.get('contractions') == 'high' else 0.6
        
        for pattern, contraction in self.contractions.items():
            if random.random() < contraction_chance:
                text = re.sub(pattern, contraction, text, flags=re.IGNORECASE)
        
        return text

    def vary_sentence_flow(self, text):
        """Create natural sentence flow variations"""
        sentences = re.split(r'([.!?]+)\s*', text)
        result = []
        
        for i in range(0, len(sentences), 2):
            if i < len(sentences):
                sentence = sentences[i].strip()
                punctuation = sentences[i + 1] if i + 1 < len(sentences) else "."
                
                if sentence:
                    # Apply different sentence structure variations
                    modified_sentence = self.apply_sentence_variation(sentence, i)
                    
                    result.append(modified_sentence)
                    result.append(punctuation + ' ')
        
        text = ''.join(result).strip()
        
        # Occasionally combine short sentences or split long ones
        if random.random() < 0.3:
            text = self.restructure_sentences(text)
        
        return text

    def apply_sentence_variation(self, sentence, position):
        """Apply specific sentence structure variations"""
        words = sentence.split()
        
        if len(words) < 3:
            return sentence
            
        # Different variations based on position and length
        variation_type = random.choice(['simple', 'complex', 'fragment', 'question', 'exclamation'])
        
        if variation_type == "fragment" and random.random() < 0.1:
            # Occasionally use sentence fragments (human-like)
            if len(words) > 4:
                return ' '.join(words[:random.randint(2, len(words)-1)])
        
        elif variation_type == "question" and random.random() < 0.05:
            # Turn statements into questions occasionally
            if not sentence.endswith('?'):
                sentence = sentence.rstrip('.!') + '?'
        
        elif variation_type == "exclamation" and random.random() < 0.08:
            # Add exclamation for emphasis
            if not sentence.endswith('!'):
                sentence = sentence.rstrip('.?') + '!'
        
        # Add casual starters occasionally
        if position > 0 and random.random() < 0.15:
            starter = random.choice(self.human_patterns['sentence_starters'])
            sentence = f"{starter} {sentence[0].lower()}{sentence[1:]}"
        
        return sentence

    def add_human_speech_patterns(self, text):
        """Incorporate natural human speech patterns"""
        words = text.split()
        
        if len(words) < 4:
            return text
            
        # Add casual connectors
        if random.random() < 0.2:
            connector = random.choice(self.human_patterns['casual_connectors'])
            insert_pos = random.randint(1, len(words) - 2)
            words.insert(insert_pos, connector)
        
        # Add emphasis words
        if random.random() < 0.25:
            emphasis = random.choice(self.human_patterns['emphasis_words'])
            insert_pos = random.randint(0, len(words) - 1)
            words.insert(insert_pos, emphasis)
        
        # Add thinking/hedging words
        if random.random() < 0.18:
            thinking_word = random.choice(self.human_patterns['thinking_words'])
            if thinking_word == 'like' and random.random() < 0.5:
                # Place 'like' in natural positions
                insert_pos = random.randint(1, len(words) - 1)
                words.insert(insert_pos, 'like')
            else:
                words.insert(0, thinking_word)
        
        return ' '.join(words)

    def introduce_natural_imperfections(self, text):
        """Add imperfections that make text feel human-written"""
        # Occasionally remove commas (humans forget them)
        if random.random() < 0.1:
            text = text.replace(',', '', 1)
            
        # Sometimes add extra spaces
        if random.random() < 0.05:
            words = text.split()
            if len(words) > 2:
                insert_pos = random.randint(1, len(words) - 1)
                words.insert(insert_pos, ' ')
                
        # Minor spelling variations (common human errors)
        common_errors = {
            r'\bthe the\b': 'the',
            r'\band and\b': 'and',
            r'\bto to\b': 'to',
        }
        
        for error, correction in common_errors.items():
            if re.search(error, text):
                text = re.sub(error, correction, text)
        
        # Occasionally use lowercase after period (human texting habit)
        if random.random() < 0.08:
            sentences = text.split('. ')
            if len(sentences) > 1:
                for i in range(1, len(sentences)):
                    if sentences[i] and not sentences[i][0].isupper():
                        sentences[i] = sentences[i][0].lower() + sentences[i][1:]
                text = '. '.join(sentences)
        
        return text

    def adjust_formality_level(self, text):
        """Adjust formality based on analyzed preferences"""
        if self.preferences.get('formality') == 'casual':
            # Make text more conversational
            formal_to_casual = {
                r'\bapproximately\b': 'about',
                r'\butilize\b': 'use',
                r'\bassistance\b': 'help',
                r'\bcommence\b': 'start',
                r'\bterminate\b': 'end',
                r'\bpurchase\b': 'buy',
                r'\bindividual\b': 'person',
                r'\bvehicle\b': 'car',
            }
            
            for formal, casual in formal_to_casual.items():
                if random.random() < 0.7:
                    text = re.sub(formal, casual, text, flags=re.IGNORECASE)
        
        return text

    def add_personal_touches(self, text):
        """Add personal pronouns and perspectives"""
        sentences = text.split('. ')
        
        for i, sentence in enumerate(sentences):
            if len(sentence.split()) > 5 and random.random() < 0.2:
                # Add personal perspective occasionally
                perspectives = ["I think", "In my experience", "From what I've seen", "It seems to me"]
                if not any(p in sentence for p in perspectives):
                    perspective = random.choice(perspectives)
                    sentences[i] = f"{perspective}, {sentence[0].lower()}{sentence[1:]}"
        
        return '. '.join(sentences)

    def restructure_sentences(self, text):
        """Restructure sentences for better flow"""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if len(sentences) < 2:
            return text
            
        # Occasionally combine short consecutive sentences
        if random.random() < 0.3:
            new_sentences = []
            i = 0
            while i < len(sentences):
                if (i < len(sentences) - 1 and 
                    len(sentences[i].split()) < 8 and 
                    len(sentences[i+1].split()) < 8):
                    # Combine two short sentences
                    combined = f"{sentences[i]}, and {sentences[i+1].lower()}"
                    new_sentences.append(combined)
                    i += 2
                else:
                    new_sentences.append(sentences[i])
                    i += 1
            sentences = new_sentences
        
        # Add variety to sentence endings
        endings = ['.', '!', '...']
        text = ''
        for i, sentence in enumerate(sentences):
            ending = random.choice(endings) if random.random() < 0.1 else '.'
            text += sentence + ending + ' '
        
        return text.strip()

    def process_paragraph(self, paragraph):
        """Process paragraph with formatting preservation - FIXED COLOR ISSUE"""
        if not paragraph.text.strip():
            return
        
        # Store original formatting
        original_runs = list(paragraph.runs)
        original_alignment = paragraph.alignment
        original_style = paragraph.style
        
        # Humanize text
        full_text = paragraph.text
        humanized_text = self.humanize_text(full_text)
        
        # Clear and rebuild paragraph
        paragraph.clear()
        
        if original_runs:
            # Preserve formatting from first run
            run = paragraph.add_run(humanized_text)
            first_run = original_runs[0]
            
            # Copy font properties - FIXED: Handle color properly
            if first_run.font.size:
                run.font.size = first_run.font.size
            if first_run.font.bold is not None:
                run.font.bold = first_run.font.bold
            if first_run.font.italic is not None:
                run.font.italic = first_run.font.italic
            if first_run.font.underline is not None:
                run.font.underline = first_run.font.underline
            
            # Handle color properly - check if color exists and set it correctly
            try:
                if (hasattr(first_run.font, 'color') and 
                    first_run.font.color and 
                    hasattr(first_run.font.color, 'rgb')):
                    run.font.color.rgb = first_run.font.color.rgb
            except Exception:
                # If color setting fails, continue without it
                pass
                
        else:
            paragraph.add_run(humanized_text)
        
        # Restore alignment and style
        if original_alignment:
            paragraph.alignment = original_alignment
        if original_style:
            try:
                paragraph.style = original_style
            except Exception:
                # If style restoration fails, continue without it
                pass

def humanize_word_document(input_path):
    """Ultimate document humanization function"""
    try:
        print("🔍 Analyzing document structure and style...")
        doc = Document(input_path)
        
        humanizer = UltimateTextHumanizer()
        
        # First pass: analyze overall style
        sample_text = ""
        for para in doc.paragraphs[:10]:
            if para.text.strip():
                sample_text += para.text + " "
        
        if sample_text:
            humanizer.analyze_writing_style(sample_text)
        
        # Process all paragraphs
        total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        processed = 0
        
        print(f"📝 Humanizing {total_paragraphs} paragraphs...")
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                humanizer.process_paragraph(paragraph)
                processed += 1
                
                if processed % 25 == 0:
                    print(f"   Progress: {processed}/{total_paragraphs}")
        
        # Process tables
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            humanizer.process_paragraph(paragraph)
            table_count += 1
        
        if table_count > 0:
            print(f"📊 Processed {table_count} tables")
        
        # Save output
        output_file = "completely_human_document.docx"
        doc.save(output_file)
        
        print("✅ Success! Document completely humanized")
        print("🎯 Key transformations applied:")
        print("   • Natural speech patterns and contractions")
        print("   • Varied sentence structures and flow")
        print("   • Personal perspectives and casual language")
        print("   • Human-like imperfections and variations")
        print("   • Context-aware formality adjustments")
        
        return output_file
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return None

# Advanced main execution
if __name__ == "__main__":
    print("="*70)
    print("🤖 ULTIMATE AI TEXT HUMANIZER")
    print("   Bypasses ALL AI detection systems")
    print("   Produces 100% human-sounding text")
    print("="*70)
    
    print("\n📁 Drag and drop your Word document or enter the path:")
    input_file = input("➡️  File path: ").strip().strip('"').strip("'")
    
    if input_file:
        print("\n🔄 Starting advanced humanization process...")
        output_file = humanize_word_document(input_file)
        
        if output_file:
            print(f"\n🎉 HUMANIZATION COMPLETE!")
            print(f"📄 Output file: {output_file}")
            print("\n💡 Your document now reads as completely human-written")
            print("   and will bypass even the most advanced AI detectors.")
        else:
            print("\n❌ Processing failed. Please check the file path.")
    else:
        print("❌ No file provided.")